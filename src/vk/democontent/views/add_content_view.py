# -*- coding: utf-8 -*-

# from vk.democontent import _
from Products.Five.browser import BrowserView
from zope.interface import implementer
from zope.interface import Interface
from plone import api
import lorem
from lorem.text import TextLorem
from plone.app.textfield import RichTextValue
import pkg_resources
from random import choice, choices, randrange
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from io import BytesIO
from plone.namedfile import NamedBlobFile, NamedBlobImage
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from docx import Document
import xlsxwriter
from PIL import Image, ImageDraw


# from Products.Five.browser.pagetemplatefile import ViewPageTemplateFile


class IAddContentView(Interface):
    """Marker Interface for IAddContentView"""


@implementer(IAddContentView)
class AddContentView(BrowserView):
    # If you want to define a template here, please remove the template from
    # the configure.zcml registration of this view.
    # template = ViewPageTemplateFile('add_content_view.pt')

    def __call__(self):
        # Implement your own actions:
#        portal = api.portal.get()
#        obj = api.content.create(type='Document',title='My Content', container=portal)

        # Folderliste zur Platzierung der Inhalte (standardmäßig nur der aktuelle Ordner)
        self.folder_list = [self.context,]

        # Wenn das Formular gesendet wurde
        if self.request.method == 'POST':
            self.number_of_documents = int(self.request.get('number_of_documents', ''))
            self.number_of_news_items = int(self.request.get('number_of_news_items', ''))
            self.number_of_links = int(self.request.get('number_of_links', ''))
            self.number_of_pdf = int(self.request.get('number_of_pdf', ''))
            self.number_of_docx = int(self.request.get('number_of_docx', ''))
            self.number_of_xlsx = int(self.request.get('number_of_xlsx', ''))
            self.number_of_jpeg = int(self.request.get('number_of_jpeg', ''))
            self.number_of_folders = int(self.request.get('number_of_folders', ''))
            self.max_folder_depth = int(self.request.get('max_folder_depth', ''))


            self.create_folders()
            self.create_contents(self.number_of_documents, 'Document')
            self.create_contents(self.number_of_links, 'News Item')
            self.create_contents(self.number_of_links, 'Link')
            self.create_contents( self.number_of_pdf, 'File', 'pdf')
            self.create_contents( self.number_of_docx, 'File', 'docx')
            self.create_contents( self.number_of_xlsx, 'File', 'xlsx')
            self.create_contents( self.number_of_docx, 'Image')


        return self.index()


    def create_folders(self):

        while len(self.folder_list) - 1 < self.number_of_folders:
            # select random container in max folder_depth
            container = choice(self.folder_list)
            # ist maximale Tiefe erreicht?
            depth = len(container.getPhysicalPath()) - len(self.context.getPhysicalPath())
            while depth > self.max_folder_depth:
                # dann anderen Folder wählen
                container = choice(self.folder_list)
                depth = len(container.getPhysicalPath()) - len(self.context.getPhysicalPath())
            # Folder erzeugen
            folder = self.create_content('Folder', container = container)
            # an Liste anhängen
            self.folder_list.append(folder)


    def get_random_words(self,count=10):
        file_path = pkg_resources.resource_filename('vk.democontent.views', 'wortliste.txt')
        with open(file_path, 'r', encoding='utf-8') as file:
            words = file.readlines()
            words = [element.strip() for element in words]
        selection = choices(words, k = count)
        return selection

    def create_contents(self, number, type, subtype = None):
        for i in range(number):
            container = choice(self.folder_list)
            self.create_content(type, subtype, container)

    def create_richtextvalue(self, words_to_use):
        return RichTextValue(raw='<p>'+ TextLorem(psep='</p> \n\n <p>', words = words_to_use).text() +' </p>', mimeType ='text/html', outputMimeType = 'text/x-html-safe')

    def create_file(self, buffer, suffix, content_type):
        buffer.seek(0)
        file_data = buffer.read()
        file_name = self.get_random_words(1)[0]+"."+ suffix
        return NamedBlobFile(data=file_data, filename=file_name, contentType= content_type)

    def create_file_pdf(self, words_to_use):
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        doc.title = TextLorem(srange=(1,5),ssep='', words = words_to_use).sentence()[:-1] # ohne Punkt
        story = []
        text = TextLorem(srange=(2,8),
            prange=(10,30),
            trange=(60,100),
            words=words_to_use).text()

        # Text in Absätze aufteilen (nach doppelten Zeilenumbrüchen)
        paragraphs = text.split('\n\n')

        # Jeden Absatz in ein Paragraph-Objekt umwandeln und zur Story hinzufügen
        for paragraph_text in paragraphs:
            p = Paragraph(paragraph_text.strip())
            story.append(p)
            # Einen Spacer (Leerzeile) hinzufügen, um die Absätze zu trennen
            story.append(Spacer(1, 12))  # 12 Punkte Abstand zwischen den Absätzen
        doc.build(story)
        return self.create_file(buffer, 'pdf', 'application/pdf' )

    def create_file_docx(self, words_to_use):
        buffer = BytesIO()
        doc = Document()
        doc.add_heading(TextLorem(srange=(1,5),ssep='', words = words_to_use).sentence()[:-1])
        text = TextLorem(srange=(2,8),
            prange=(10,30),
            trange=(60,100),
            words=words_to_use).text()

        # Text in Absätze aufteilen (nach doppelten Zeilenumbrüchen)
        paragraphs = text.split('\n\n')
        # Jeden Absatz hinzufügen
        for paragraph_text in paragraphs:
            doc.add_paragraph(paragraph_text.strip())
        doc.save(buffer)
        return self.create_file(buffer, 'docx', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' )

    def create_file_xlsx(self, words_to_use):
        buffer = BytesIO()
        workbook = xlsxwriter.Workbook(buffer)
        worksheet = workbook.add_worksheet()
        text = TextLorem(words = words_to_use).paragraph()
        worksheet.write('A1', text)
        workbook.close()
        return self.create_file(buffer, 'xlsx', 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet' )

    def create_image(self):
        # Größe des Bildes
        width, height = 800, 600
        # Hintergrundfarbe festlegen (weiß)
        background_color = (randrange(255), randrange(255), randrange(255))
        # Textfarbe festlegen (schwarz)
        text_color = (0, 0, 0)
        # Eine neue Image-Instanz erstellen
        image = Image.new("RGB", (width, height), background_color)
        # Eine Zeichenfläche erstellen
        draw = ImageDraw.Draw(image)
        # Position und Text festlegen
        text = "Hallo, Welt!"
        text_position = (width // 4, height // 3)
        # Text auf die Zeichenfläche zeichnen
        draw.text(text_position, text, fill=text_color)

        buffer = BytesIO()
        image.save(buffer,'JPEG')

        buffer.seek(0)
        file_data = buffer.read()
        file_name = self.get_random_words(1)[0]+".jpg"
        return NamedBlobImage(data=file_data, contentType = "image/jpeg", filename=file_name)

    def create_content(self, type, subtype = None, container = None):
            words_to_use = self.get_random_words(4000)
            d = {}
            d['type'] = type
            d['title'] = TextLorem(srange=(1,5),ssep='', words = words_to_use).sentence()[:-1] # ohne Punkt
            d['description'] = TextLorem(words = words_to_use).paragraph()
            if container == None:
                d['container'] = self.context
            else:
                d['container'] = container

            if type == 'Document':
                d['text'] = self.create_richtextvalue(words_to_use)

            if type == 'File':

                if subtype == "pdf":
                    d['file'] = self.create_file_pdf(words_to_use)

                if subtype == "docx":
                    d['file'] = self.create_file_docx(words_to_use)

                if subtype == "xlsx":
                    d['file'] = self.create_file_xlsx(words_to_use)

            if type == 'Image':
                d['image'] = self.create_image()

            if type == 'Link':
                d['remoteUrl'] = "https://plone.org"

            if type == 'News Item':
                d['text'] = self.create_richtextvalue(words_to_use)
                d['image'] = self.create_image()

            if type == 'Folder':
                pass

            obj = api.content.create(**d)

            return obj

